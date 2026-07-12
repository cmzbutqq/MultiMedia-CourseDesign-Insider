from __future__ import annotations

import hashlib
import re
import shutil
import subprocess
import tempfile
import textwrap
from pathlib import Path
from typing import Iterable

from PIL import Image
from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.opc.constants import RELATIONSHIP_TYPE
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
DOCS_DIR = ROOT / "docs"
TEMPLATE_PATH = DOCS_DIR / "报告模板.docx"
OUTPUT_PATH = DOCS_DIR / "多媒体课程作业实验报告.docx"

PDFLATEX = "/Library/TeX/texbin/pdflatex"
LATEX = "/Library/TeX/texbin/latex"
DVIPNG = "/Library/TeX/texbin/dvipng"
PDFTOPPM = (
    "/Users/cmz/.cache/codex-runtimes/codex-primary-runtime/dependencies/bin/pdftoppm"
)

TITLE = "基于 WebGL2 的黑洞与多天体可视化系统设计与实现"
BODY_FONT = "SimSun"
HEADING_FONT = "SimHei"
LATIN_FONT = "Times New Roman"
CODE_FONT = "Courier New"
FORMULA_FONT_PT = 11
FORMULA_RASTER_DPI = 600
FIGURE_RASTER_DPI = 600

FIGURES = {
    "single": ROOT / "pics/1-1-single-entity.png",
    "double": ROOT / "pics/1-2-double-entity.png",
    "kepler": ROOT / "pics/3-1-Kepler.png",
    "nbody": ROOT / "pics/3-2-N-body.png",
    "dt": ROOT / "pics/3-3-N-body-Time-Step.png",
    "gesture": ROOT / "pics/2-3-gesture-control.png",
    "camera": ROOT / "pics/2-2-camera-rotate.png",
    "fsr_off": ROOT / "pics/4-2-no-FSR.png",
    "fsr_on": ROOT / "pics/4-3-FSR.png",
    "aa": ROOT / "pics/4-6-anti-aliasing.png",
    "skybox": ROOT / "pics/4-9-skybox4.png",
    "doppler": ROOT / "pics/5-2-doppler shift.png",
    "beaming": ROOT / "pics/5-3-beam intensification.png",
    "bloom": ROOT / "pics/5-4-boom-layers-low.png",
    "recording": ROOT / "pics/6-recording.png",
    "audio": ROOT / "pics/7-audio.png",
    "display": ROOT / "pics/9-display.png",
    "gh_actions": ROOT / "pics/GitHub Actions.png",
    "gh_pages": ROOT / "pics/GitHub pages.png",
}


SECTIONS = [
    {
        "title": "一、项目资源与系统概览",
        "opening_resources": True,
        "paragraphs": [
            "项目主体位于 web/，采用 WebGL2、TypeScript 与 Vite 组织实时渲染和交互；server/ 提供可选的服务端手势识别；根目录 docker-compose.yml、README.md 与 .github/workflows/ 共同承担构建、测试、部署和交付说明。核心控制器是 web/src/main.ts，它在同一帧循环里统一推进物理演化、着色器参数、后处理通道、GUI、录制回放、音频和输入状态。",
            "场景状态由 SceneState 描述，核心字段包括活跃天体数、天体数组、动力学模式、中心引力常数、N 体引力常数、软化半径、时间步长与时间扭曲参数。项目当前最多支持 5 个活跃天体，类型分为黑洞、白洞和中子星；预设分为单天体、双天体、开普勒演示与 N 体演示。统一状态结构使渲染、物理、交互、录制和音频都围绕同一份数据工作，因此参数修改、录制回放和场景预设之间可以直接互通。",
            "系统的主链路可以概括为“状态推进 - 光线着色 - 显示重建 - 交互反馈”。shader 层负责黑洞透镜、吸积盘、相对论视效和天空盒；physics.ts 负责多体运动和时间扭曲近似；camera、handGesture、recordingManager 与 ambientAudio 分别承担视角、手势、可复现实验和参数声化；Docker、CI、CD 与 GitHub Pages 则把本地工程进一步转换为可复现、可验证、可公开演示的成果。",
        ],
        "figures": [
            ("single", "图 1  单黑洞场景、吸积盘、天空盒和参数面板的整体界面。", 5.9),
        ],
    },
    {
        "title": "二、需求分析与系统设计思路",
        "paragraphs": [
            "本项目属于作品设计型多媒体课程实践，目标不是只生成若干静态效果截图，而是构建一个可实时交互、可参数化控制、可录制回放、可公开演示的黑洞与多天体可视化系统。系统既要能够展示黑洞、白洞、中子星和多天体轨道等核心内容，也要支持课堂演示、结果复现和后续扩展，因此需求分析必须同时覆盖功能完整性、演示表现力和工程可交付性。",
            "功能需求可以概括为六组：其一，完成黑洞、白洞和中子星等天体的实时渲染，并支持吸积盘、相对论视觉效果和天空盒背景；其二，支持开普勒与 N 体两类动力学模式，以及轨迹线、场景预设和时间扭曲等演示能力；其三，提供鼠标与摄像头手势两类视角操控方式；其四，提供抗锯齿、上采样、渲染倍率、后处理和显示配置等画面调节能力；其五，支持脚本镜头、参数 JSON 录制与回放，以及 Web Audio 氛围音频联动；其六，支持容器化、自动化测试和 GitHub Pages 演示部署，保证作品能够稳定交付。",
            "非功能需求主要包括实时性、可解释性、可演示性和可维护性。实时性要求系统在浏览器环境中完成主要渲染与交互；可解释性要求参数、场景状态和视觉效果之间具有清晰对应关系，便于课程展示和教师审阅时说明；可演示性要求系统具备预设、录制回放、截图和公开视频页面等能力；可维护性要求代码结构清晰，前后端边界明确，并具备 Docker、CI/CD 等工程化支撑。",
            "系统设计思路采用“统一状态驱动 + 模块化分层实现”的方案。SceneState 作为核心状态中心，统一承载天体、动力学、显示、录制和音频相关信息；main.ts 作为调度入口，在每一帧中协调物理推进、渲染更新、后处理、GUI、录制回放和输入响应；shader 层负责黑洞透镜、吸积盘、后处理与显示重建；physics.ts 负责多体演化与时间扭曲近似；camera、handGesture、recordingManager 和 ambientAudio 分别负责视角、手势、实验复现和参数声化；server/ 则提供可选的服务端识别能力。这样形成的设计既能支撑课程展示，也便于按模块分工开发和后续扩展。",
        ],
    },
    {
        "title": "三、小组组员分工说明",
        "blocks": [
            {
                "type": "paragraph",
                "text": "项目采用分模块协同推进的方式完成，分工以“核心主链路 + 专项功能模块 + 展示与文档支撑”为主。下面按照成员分别说明承担的功能模块、实现内容和参与方式。"
            },
            {
                "type": "table",
                "caption": "表 1  小组成员分工说明。",
                "headers": ["成员", "分工说明"],
                "rows": [
                    ["陈梦泽", "陈梦泽承担了整体推进和主要开发工作，是项目落地的核心组织者与实现者。他负责了网页端黑洞可视化系统的主体搭建与持续完善，包括前期框架建立、后期界面交互优化、场景效果调整、性能改进以及文档与展示材料整理等内容。同时，他还承担了各部分功能的整合衔接工作，推动项目从分散功能逐步形成完整可演示的系统。除代码开发外，演示视频由陈梦泽负责录制，实验报告也由他统稿撰写，对项目最终成果展示起到了主导作用。"],
                    ["莫镇嘉", "莫镇嘉主要负责手势交互相关功能的设计与实现，为项目增加了更直观、更有展示性的交互方式。他完成了手势识别服务端、前后端通信、相关部署配置以及配套说明文档，使这一部分不仅能运行，而且便于集成和展示。除此之外，他还参与了抗锯齿、HTTPS 支持等体验优化工作，并为实验报告提供了截图素材，对项目展示效果的提升帮助很大。"],
                    ["廖晨扬", "廖晨扬主要负责多天体场景和轨道运动相关功能的实现，扩展了项目的模拟内容和表现层次。他完成了多天体管理、轨迹显示以及相关运动逻辑的设计，使系统不再局限于单一黑洞展示，而能够呈现更丰富的天体运动过程。这部分工作增强了项目的完整性，也提升了演示时的观感和内容深度。演示视频的剪辑和字幕制作也有他的重要参与。"],
                    ["陈紫暄", "陈紫暄主要负责录制回放、时间变化效果和音频氛围相关功能的补充，让项目更适合展示和演示。她实现了镜头与参数的录制回放功能，使演示内容能够更稳定、更连贯地复现；同时也完善了环境音效和时间扭曲效果，增强了作品的沉浸感和表现力。演示视频的剪辑和字幕制作也有她的参与。"],
                    ["曹丹", "曹丹主要在相对论视觉效果和理论支撑方面提供了贡献。她参与了相关效果的验证与调整，使部分画面表现更加贴近预期的物理现象。同时，在实验报告撰写过程中，她负责查找和补充理论材料，为报告内容的完整性和科学性提供了支持。"],
                ],
            },
        ],
    },
    {
        "title": "四、技术栈选型",
        "blocks": [
            {
                "type": "paragraph",
                "text": "技术栈选型围绕三个约束展开：核心效果必须在浏览器里实时运行，演示页面必须能直接公开部署，摄像头手势既要支持纯前端路径，也要支持可独立调试的服务端路径。项目据此采用原生图形接口与轻量工程组织，把渲染、交互、多媒体、识别和交付拆成几层边界清楚的能力组合。"
            },
            {
                "type": "table",
                "caption": "表 2  技术栈分层与选型依据。",
                "headers": ["层级", "实际技术", "直接职责", "选型依据"],
                "rows": [
                    ["浏览器渲染", "WebGL2 + GLSL", "全屏光线步进、浮点离屏缓冲、后处理 pass", "直接控制 shader、FBO 和采样状态，同时保持浏览器端可部署性"],
                    ["工程语言与构建", "TypeScript + Vite", "状态类型约束、ESM 开发、raw shader 导入、静态构建", "状态流跨 GUI、录制、uniform 与 JSON，轻量构建比重型框架更贴合项目形态"],
                    ["交互与多媒体", "lil-gui + pointer events + Web Audio API", "参数面板、轨道相机、实时音频映射", "控制链路短，适合实验型界面与连续参数调节"],
                    ["手势识别", "MediaPipe Hands；Flask + MediaPipe + OpenCV", "本地低延迟识别；服务端隔离推理", "同时兼顾公开演示、本地响应和后端独立调试"],
                    ["交付", "Docker + Nginx + GitHub Actions + GitHub Pages", "环境收敛、统一入口、自动验证、静态部署", "前端适合直接静态发布，服务端依赖更适合容器封装"],
                ],
            },
            {
                "type": "paragraph",
                "text": "从负载结构看，最重的开销仍然是片元阶段的步进和后处理，因此主要计算工作集中在 WebGL2 片元管线；手势部分同时保留本地和服务端两条路径，对应的是延迟、隔离性和部署复杂度之间的工程取舍。"
            },
            {
                "type": "formula",
                "latex": r"""\begin{aligned}
T_{\mathrm{frame}} &\approx T_{\mathrm{CPU}} + W_rH_r\Bigl(C_{\mathrm{trace}}(\bar N_{\mathrm{step}},N_{\mathrm{body}})+C_{\mathrm{post}}\Bigr)+T_{\mathrm{present}},\\
C_{\mathrm{trace}} &\sim O\!\left(\bar N_{\mathrm{step}}N_{\mathrm{body}}\right)
\end{aligned}""",
            },
            {
                "type": "paragraph",
                "text": "手势识别链路的延迟来源分成两类。本地链路主要由摄像头采集、关键点推理和平滑构成；服务端链路还会额外叠加图像编码、网络传输、后端排队与回传开销，因此它更强调依赖隔离、日志记录和独立调试。"
            },
            {
                "type": "paragraph",
                "text": "静态构建与公开部署采用如下工程结构："
            },
            {
                "type": "code",
                "code": """
web/
  src/        # TypeScript orchestration
  shader/     # GLSL passes
  public/     # skybox and static textures
  dist/       # generated by `vite build`, deployed to GitHub Pages

server/
  server.py
  requirements.txt
                """,
            },
        ],
    },
    {
        "title": "五、黑洞场景实时渲染",
        "paragraphs": [
            "黑洞主体采用“全屏四边形 + 片元内光线步进”的实时方案。blackhole_main.frag 为每个片元生成一条从 cameraWorld 出发的视线，在 traceColor() 循环中反复推进位置和方向，并在循环内部累积吸积盘体发光与天体命中结果；当射线离开主要相互作用区后，再用最终 rayDir 对天空盒取样。这样的结构把连续光路离散成有限次局部更新，适合浏览器端实时渲染。",
            "黑洞附近的光子传播受零测地线约束。完整理论里，光线沿 Kerr 时空传播，常用的守恒量是光子能量 $E=-p_t$、轴向角动量 $L_z=p_\\phi$ 和 Carter 常数 $Q$；它们共同决定光线会不会穿过赤道面、会不会多次绕转，以及最终落入视界还是逃向无穷远。事件视界、能层、光子球和稳定圆轨道边界分别控制不可逃逸区、拖拽旋转区、阴影边界和吸积盘内缘。若限制在 Schwarzschild 球对称情形，黑洞阴影边缘可以追溯到逃逸轨道与俘获轨道之间的临界撞击参数 $b_{\\mathrm{crit}}=3\\sqrt{3}M$：当远处入射光线的撞击参数小于这个阈值时，光线会被黑洞捕获；大于它时则仍有机会逃向无穷远。",
            "黑洞图像里还要把事件视界、光子球、阴影和光子环分开。事件视界是不可逃逸边界；Schwarzschild 黑洞的光子球位于 $r=3M$；远处观察者看到的阴影半径由临界撞击参数决定，对应成像平面上的 $r_{\\mathrm{shadow}}=3\\sqrt{3}M=\\frac{3\\sqrt{3}}{2}r_s$；更细的高亮光子环则来自光线在接近光子球时经历一次或多次绕转后再逃逸。把这几个概念拆开，画面中的黑色核心、亮环和背景畸变就不容易混淆。",
            "实时渲染阶段采用 Schwarzschild 情形的 Binet 偏折核向量近似来更新背景光线方向；自旋参数主要作用于吸积盘内缘位置和盘面局部轨道速度，从而把背景透镜、盘体结构和自旋视效稳定组织在同一条片元链路中。",
            "当前渲染实现区分了两套半径。bodySize 对应天体核心命中与内部着色，bodySurfaceRadius 则用于估计动态步长和追踪范围；对黑洞而言，后者会按视觉尺寸放大，使光线在主体附近自动细分。这样的分工把“可见核心”和“数值采样安全边界”拆开处理，既保留黑洞中心的锐利轮廓，也避免过远区域浪费步数。",
            "吸积盘的理论基础可以从薄盘模型理解：角动量向外输运、引力势能向热辐射转化，盘面通量和温度随半径变化；完整图像还应沿观测光线积分辐射转移方程。实时实现把盘体建模为围绕 body0 的体发光介质，并用半径衰减、竖向密度、分形噪声、色图映射和相对论亮度修正组织出“内缘更亮、外缘更冷、纹理受湍流破坏”的总体观感。",
            "当前盘体密度由“几何包络 × 径向幂律 × 多尺度乘性噪声”组成。代码先用外包络和内缘 smoothstep 裁出有限厚度盘层，再在类球坐标变量上叠乘多级 simplex noise，并让奇偶层沿相反角向漂移。这样的乘性分形模型不是简单给盘面贴一张噪声图，而是在用极少参数模拟剪切流、局部团块和旋臂状不均匀结构，使亮纹既有方向性又不会过于规则。",
            "黑洞成像对应的不只是单条光线，而是一个有限像素张角内的测地线束。束宽、剪切和会聚由测地线偏差方程控制，它们决定了高亮细丝会不会在屏幕上跳闪、环像会不会断裂；靠近 $r=3M$ 的临界轨道时，这种束传播还会把多次绕转的光线压成更细的高亮光子环。当前实时链路采用单射线近似，并配合动态步长、时域融合和屏幕空间后处理来稳定画面。",
            "盘面纹理中的噪声也有明确的数学来源。shader 使用的是 3D simplex gradient noise：先把欧氏网格斜切到 simplex 单纯形晶格，再对 4 个角点的梯度投影做径向衰减求和。与规则立方格上的 value noise 或普通 Perlin noise 相比，simplex 噪声采样点更少、方向伪影更弱，更适合在浏览器 shader 中反复迭代生成吸积盘湍流纹理。",
            "shader 中的主实现为 body0 绑定一套全局吸积盘参数，因此盘的内边界、厚度、水平密度、垂直密度、噪声层级、旋转速度和亮度都围绕同一黑洞稳定调节。这一设计把“盘面物理参数”与“主体透镜参数”明确拆开，使黑洞主体、盘体、背景星场和后处理可以在同一条片元链路里协同工作。界面里还可以单独切换 Lensing、Show Core 和吸积盘启停，因此同一场景也适合做透镜、核心和盘体的消融观察。",
        ],
        "formulas": [
            r"""\begin{aligned}
0 &= ds^2 = g_{\mu\nu}dx^\mu dx^\nu,\\
\frac{d^2x^\mu}{d\lambda^2} + \Gamma^\mu_{\alpha\beta}\frac{dx^\alpha}{d\lambda}\frac{dx^\beta}{d\lambda} &= 0
\end{aligned}""",
            r"""\begin{aligned}
ds^2 &= -\left(1 - \frac{2M}{r}\right)dt^2 + \left(1 - \frac{2M}{r}\right)^{-1}dr^2 \\
&\quad + r^2\left(d\theta^2 + \sin^2\theta\, d\phi^2\right)
\end{aligned}""",
            r"""\begin{aligned}
ds^2 &= -\left(1-\frac{2Mr}{\Sigma}\right)dt^2 - \frac{4Mar\sin^2\theta}{\Sigma}\, dt\, d\phi + \frac{\Sigma}{\Delta}dr^2 + \Sigma\, d\theta^2 \\
&\quad + \left(r^2 + a^2 + \frac{2Ma^2r\sin^2\theta}{\Sigma}\right)\sin^2\theta\, d\phi^2, \\
\Sigma &= r^2 + a^2\cos^2\theta,\qquad \Delta = r^2 - 2Mr + a^2
\end{aligned}""",
            r"""\begin{aligned}
\mathcal{E} &= -p_t,\qquad L_z = p_\phi,\\
Q &= p_\theta^2 + \cos^2\theta\left(a^2\mathcal{E}^2 - \frac{L_z^2}{\sin^2\theta}\right),\\
\Sigma^2\left(\frac{dr}{d\lambda}\right)^2 &= \left[\mathcal{E}(r^2+a^2)-aL_z\right]^2 - \Delta\left[(L_z-a\mathcal{E})^2 + Q\right],\\
\Sigma^2\left(\frac{d\theta}{d\lambda}\right)^2 &= Q - \cos^2\theta\left(\frac{L_z^2}{\sin^2\theta} - a^2\mathcal{E}^2\right)
\end{aligned}""",
            r"""\frac{D^2\xi^\mu}{d\lambda^2}=R^\mu{}_{\nu\alpha\beta}k^\nu k^\alpha \xi^\beta""",
            r"""\begin{aligned}
r_H &= M + \sqrt{M^2-a^2},\\
r_{\mathrm{ergo}}(\theta) &= M + \sqrt{M^2-a^2\cos^2\theta},\\
r_{\mathrm{ph}}^{\mathrm{Schw}} &= 3M,\qquad
b_{\mathrm{crit}}^{\mathrm{Schw}} = 3\sqrt{3}\,M
\end{aligned}""",
            r"""\begin{aligned}
r_s &= 2M,\\
r_{\mathrm{shadow}}^{\mathrm{Schw}} &= 3\sqrt{3}\,M = \frac{3\sqrt{3}}{2}r_s
\end{aligned}""",
            r"""\begin{aligned}
\frac{d^2u}{d\phi^2} + u &= 3Mu^2,\qquad u=\frac{1}{r},\\
\bm{a}_{\mathrm{bend}} &\approx -\frac{3}{2}\,\frac{h^2\, \bm{r}}{\lVert \bm{r}\rVert^5},\qquad h^2=\lVert \bm{r}\times \bm{v}\rVert^2
\end{aligned}""",
            r"""\begin{aligned}
F(r) &= \frac{3GM\dot{M}}{8\pi r^3}\left(1-\sqrt{\frac{r_{\mathrm{in}}}{r}}\right),\\
T(r) &= \left(\frac{F(r)}{\sigma}\right)^{1/4}
\end{aligned}""",
            r"""\begin{aligned}
\frac{dI_\nu}{ds} &= j_\nu - \alpha_\nu I_\nu,\\
j(r,z,\theta,\varphi,t) &= A(r,z)\,S_{\mathrm{in}}(r)\,r^{-p_h}
\prod_{k=1}^{K}\left[0.5+0.5\,\mathcal{N}\!\left(\sigma_k r,\, m_k\theta+(-1)^k\omega t,\, n_k\varphi\right)\right],\\
A(r,z) &= \max\!\left(0,1-\left\lVert \left(\frac{x}{R_{\mathrm{out}}},\frac{z}{H},\frac{y}{R_{\mathrm{out}}}\right)\right\rVert\right)\left(1-\frac{|z|}{H}\right)^{p_v}
\end{aligned}""",
            r"""\begin{aligned}
n(\bm{v}) &= 42\sum_{i=0}^{3} m_i^4\, \bigl(\bm{p}_i\cdot \bm{x}_i\bigr),\\
m_i &= \max\!\left(0.6-\|\bm{x}_i\|^2,\,0\right)
\end{aligned}""",
        ],
        "figures": [
            ("single", "图 2  单黑洞场景中的黑洞阴影、吸积盘和背景星场透镜畸变。", 5.9),
        ],
    },
    {
        "title": "六、白洞与中子星渲染",
        "paragraphs": [
            "白洞与中子星和黑洞共用同一套 SceneBody 数据结构，但在几何解释和视觉分支上不同。白洞在理论上对应 Schwarzschild 几何最大解析延拓中的过去视界分支，即允许物质与光线由内部涌出而不允许外部落入；在实时渲染里，这种“时间反向”的最直接可见后果就是透镜偏折方向翻转。代码通过 lensSign 改变偏折符号，使背景纹理由“汇聚”转为“发散”，从而得到与黑洞相反的透镜风格。",
            "中子星对应的是没有事件视界但紧致度极高的致密天体。其更本质的理论关键词不是“发光球体”，而是紧致度、表面引力、引力红移和强场光线弯曲：即使没有视界，观测者仍可能看见超过半个球面的表面，表面辐射也会因引力红移和发射角重映射而改变亮度分布。项目把这些强场光学后果压缩为高亮球面、自发光颜色和视角相关的 limb 亮度调制，使中子星在多天体场景中保持清晰的视觉辨识度。",
            "写到内部结构时，中子星就不再只是一个发光表面，而是一个静态球对称流体平衡问题。给定状态方程 $p=p(\\rho)$ 后，内部压强和质量分布由 Tolman-Oppenheimer-Volkoff 方程决定；Buchdahl 界 $\\mathcal{C}<4/9$ 则给出各向同性致密星在不形成视界时能达到的紧致度上界。这层理论说明了中子星为什么处在“强场但无视界”的临界一侧。",
            "统一渲染框架的优点在于三类天体共享同一套轨道、相机、录制和天空盒系统，而差异只收敛到偏折符号、核心着色和表面发光模型上。这样既保留了黑洞、白洞和中子星的物理角色差异，也让多天体场景的控制逻辑保持一致。",
        ],
        "formulas": [
            r"""\begin{aligned}
U &= T-X,\qquad V=T+X,\\
UV &= \left(1-\frac{r}{2M}\right)e^{r/(2M)}
\end{aligned}""",
            r"""\begin{aligned}
z_{\mathrm{grav}} &= \left(1-\frac{2GM}{Rc^2}\right)^{-\frac12} - 1,\\
\mathcal{C} &= \frac{GM}{Rc^2} < \frac49
\end{aligned}""",
            r"""\begin{aligned}
\frac{dp}{dr} &= -\frac{G\left(\rho + p/c^2\right)\left(m + 4\pi r^3 p/c^2\right)}{r\left(r-2Gm/c^2\right)},\\
\frac{dm}{dr} &= 4\pi r^2 \rho,\qquad p=p(\rho)
\end{aligned}""",
            r"""\mathcal{C}<\frac49 \iff R>\frac98\,r_g=\frac{9GM}{4c^2}""",
            r"""\hat{\alpha} \approx \frac{4GM}{bc^2}""",
            r"""\begin{aligned}
1-\cos\alpha &= (1-\cos\psi)\left(1-\frac{r_g}{R}\right),\\
r_g &= \frac{2GM}{c^2}
\end{aligned}""",
            r"""\begin{aligned}
L_{\mathrm{ns}}(\mu) &\approx L_0 + L_g\left(0.22 + 0.78\max(0,\mu)\right),\\
\mu &= \bm{n}\cdot \bm{v}_{\mathrm{view}}
\end{aligned}""",
        ],
        "figures": [
            ("double", "图 3  双天体场景中黑洞与白洞的同屏效果，可直观看到透镜风格差异。", 5.9),
        ],
    },
    {
        "title": "七、多天体物理：开普勒、N 体、轨迹线与场景预设",
        "paragraphs": [
            "physics.ts 提供 static、kepler 和 nbody 三种动力学模式。static 模式只渲染不演化；kepler 模式固定 body[0] 为中心天体，其余天体只受中心引力；nbody 模式则让所有活跃天体彼此相互作用。三种模式共享同一份场景状态，因此同一套参数既可用于纯展示，也可直接切换为轨道和多体动力学实验。",
            "多天体在本项目中不仅体现在物理侧，也体现在统一渲染侧。main.ts 会把每个活跃天体的中心、类型、表面半径、发光参数和吸积盘增益统一打包送入 shader，blackhole_main.frag 再在光线步进中对所有活跃天体累积透镜效应、判定最近命中并按黑洞、白洞、中子星三种类型分支着色。因此画面里的“多天体”并不是单纯多画几个球，而是多套透镜、命中与发光规则同时参与同一条片元射线路径。需要单独说明的是，当前吸积盘模型仍作为围绕 body0 的全局盘源参与渲染；因此多天体共同性主要体现在透镜、命中、轨迹与发光主体的并行组织上，而不是每个天体都自动拥有独立的相对论吸积盘。",
            "多天体问题在理论上可以写成哈密顿系统。两体情形退化为经典开普勒问题，轨道由总能量和角动量决定；N 体情形则需要在每一时刻重新累积所有相互作用。代码采用软化牛顿势来避免短距离奇异，同时保留半长轴、周期、近远拱点和速度随半径变化等关键轨道特征。",
            "速度 Verlet 属于辛积分框架下的 Strang splitting。对 $H=T+V$ 的系统，它等价于“半步动量 - 整步位置 - 半步动量”：每个子步都简单，但整体仍保持相空间体积，不会像显式欧拉那样快速拉开能量漂移。对课程实验来说，这一性质很重要，因为演示画面关心的是轨道长期形状、共振和会遇，而不是几步之内的局部截断误差。",
            "N 体模式采用速度 Verlet 积分。速度 Verlet 的核心优点是时间对称和较低的长期能量漂移，因此比显式欧拉更适合长时间展示轨道结构。进入 N 体模式时，scene.ts 还会把参考天体移近原点，并减去系统质心速度，使画面更集中地呈现相互作用本身，而不是整个系统被一个常速度整体平移。",
            "模式切换时，系统还会尽量维持参数连续性，而不是简单粗暴地重置场景。kepler 切到 nbody 时，会按中心质量推导等效的 nbodyG，并对系统做平移和去质心速度处理；nbody 切回 kepler 时，则反推 gmCentral 并把中心体速度归零。这样做的意义在于：同一个实验可以在“受控两体近似”和“完整多体相互作用”之间平滑来回，而不会因为模式切换立刻失去可比较性。",
            "轨迹线采用 TrailBuffer 环形缓冲与屏幕空间 2D 叠加。每个天体保存有限个历史点，再通过 worldToScreenPx 投影到画布上用 Canvas 绘制彩色轨迹。这个方案把“动力学历史”与“主体着色”分离开来，便于比较不同预设下的轨道形态、周期差异和混沌程度。",
        ],
        "formulas": [
            r"""\begin{alignedat}{2}
H &= \sum_i \frac{\|\bm{p}_i\|^2}{2m_i}-\sum_{i<j}\frac{Gm_im_j}{r_{ij}},\qquad &
\bm{a}_{\mathrm{lens}}^{\mathrm{tot}}(\bm{x}) &= \sum_{i=1}^{N_{\mathrm{active}}} s_i\,\bm{a}_{\mathrm{lens},i}(\bm{x})
\end{alignedat}""",
            r"""\begin{alignedat}{2}
T^2 &= \frac{4\pi^2}{GM}a^3,\qquad &
v^2 &= \mu\left(\frac{2}{r}-\frac{1}{a}\right)
\end{alignedat}""",
            r"""\begin{alignedat}{2}
\Phi_\varepsilon(r) &= -\frac{GM}{\sqrt{r^2+\varepsilon^2}},\qquad &
\bm{a}_{\mathrm{kepler}} &= -\mu\, \frac{\bm{r}}{\left(\lVert \bm{r}\rVert^2 + \varepsilon^2\right)^{3/2}}
\end{alignedat}""",
            r"""\bm{a}_i = G \sum_{j\ne i} m_j\, \frac{\bm{r}_j-\bm{r}_i}{\left(\lVert \bm{r}_j-\bm{r}_i\rVert^2+\varepsilon^2\right)^{3/2}}""",
            r"""\begin{alignedat}{2}
\bm{R}_{\mathrm{cm}} &= \frac{\sum_i m_i \bm{r}_i}{\sum_i m_i},\qquad &
\bm{V}_{\mathrm{cm}} &= \frac{\sum_i m_i \bm{v}_i}{\sum_i m_i}
\end{alignedat}""",
            r"""\begin{aligned}
\bm{v}_{t+\frac12} &= \bm{v}_t + \frac12 \bm{a}_t \Delta t,\\
\bm{r}_{t+1} &= \bm{r}_t + \bm{v}_{t+\frac12}\Delta t,\\
\bm{v}_{t+1} &= \bm{v}_{t+\frac12} + \frac12 \bm{a}_{t+1}\Delta t
\end{aligned}""",
            r"""\det\!\left(\frac{\partial(\bm{q}_{n+1},\bm{p}_{n+1})}{\partial(\bm{q}_n,\bm{p}_n)}\right)=1""",
        ],
        "figures": [
            ("kepler", "图 4  开普勒模式下的椭圆轨道演示。", 5.8),
            ("nbody", "图 5  N 体模式下多天体的相互吸引与轨迹叠加。", 5.8),
            ("dt", "图 6  时间步长改变后，N 体演化的节奏和稳定性会同步变化。", 5.8),
        ],
    },
    {
        "title": "八、高级相对论视效：多普勒偏移、束射增强与自旋近似",
        "paragraphs": [
            "相对论视效主要挂在吸积盘采样阶段，包括多普勒偏移、束射增强和自旋近似三部分。它们共同来自洛伦兹变换对频率、传播方向和辐射强度的影响，因此这一节的理论核心是观测者参考系下的频移、像差和辐射变换。",
            "更底层的写法是用四速度与光子四动量定义频移因子：观测频率与发射频率之比由发射者、观测者和光子在时空中的内积决定，而单位相空间体积上的比强度满足 $I_\\nu/\\nu^3$ 不变。实时计算时，盘面局部先构造轨道速度方向，再根据采样点指向相机的真实观察方向计算近似多普勒因子。这样得到的红蓝不对称由视线与局部轨道速度夹角直接控制。",
            "束射增强里的指数并不是任意经验参数，而是局部谱形的直接后果。若发射谱在目标频带附近满足 $F_\\nu\\propto \\nu^\\alpha$，则由 $I_\\nu/\\nu^3$ 不变量可得观测亮度近似按 $D^{3+\\alpha}$ 放大；其中 3 来自立体角、到达率与频率伸缩，$\\alpha$ 则来自谱本身对频率缩放的响应。这样一来，迎向观察者的一侧不仅颜色更偏蓝，单位立体角内接收到的能量也会更集中。",
            "当前实现把多普勒因子映射为 RGB 通道增益，用直接的颜色偏移表达红移和蓝移；束射增强则按 $D^p$ 的幂次放大局部亮度，使高速迎面区域更亮。这样做等价于把“频谱偏移”和“亮度聚束”拆成两条稳定可控的渲染通道，在实时预算内保留了“迎面侧蓝移且更亮、背离侧红移且更暗”的核心视觉规律。",
            "自旋近似对应 Kerr 黑洞的拖拽趋势。理论上，自旋既会通过 $r_{\\mathrm{ISCO}}(a_*)$ 改变盘的最内稳定边界，也会通过 frame dragging 改变局部惯性系的旋转角速度 $\\Omega=-g_{t\\phi}/g_{\\phi\\phi}$。在几何单位制 $G=c=1$ 下，自旋还会改变事件视界半径和顺、逆轨道的角速度分布；实时实现把这些趋势压缩为可控的盘内边界收缩和局部角速度修正，因此能稳定呈现顺旋侧更亮、内缘更紧的视觉印象。",
        ],
        "formulas": [
            r"""\begin{aligned}
\frac{\nu_{\mathrm{obs}}}{\nu_{\mathrm{emit}}}
&= g = \frac{u^\mu_{\mathrm{obs}}k_\mu}{u^\nu_{\mathrm{emit}}k_\nu},\qquad
\gamma = \frac{1}{\sqrt{1-\beta^2}},\qquad
D = \frac{1}{\gamma\left(1-\bm{v}\cdot \bm{n}\right)}\\
\lambda_{\mathrm{obs}} &= \frac{\lambda_{\mathrm{emit}}}{D},\qquad
\cos\theta'=\frac{\cos\theta-\beta}{1-\beta\cos\theta}
\end{aligned}""",
            r"""\begin{aligned}
\frac{I_\nu}{\nu^3} &= \text{invariant},\qquad
F_{\nu,\mathrm{emit}} \propto \nu^\alpha,\\
I_{\nu,\mathrm{obs}} &= D^3 I_{\nu,\mathrm{emit}}(\nu/D),\qquad
F_{\nu,\mathrm{obs}} \propto D^{3+\alpha}F_{\nu,\mathrm{emit}}
\end{aligned}""",
            r"""\Omega_{\mathrm{drag}} = -\frac{g_{t\phi}}{g_{\phi\phi}}""",
            r"""\begin{aligned}
r_{\mathrm{ISCO}} &= M\left[3+Z_2-\operatorname{sgn}(a_*)\sqrt{(3-Z_1)(3+Z_1+2Z_2)}\right],\\
Z_1 &= 1+(1-a_*^2)^{1/3}\left[(1+a_*)^{1/3}+(1-a_*)^{1/3}\right],\\
Z_2 &= \sqrt{3a_*^2 + Z_1^2}
\end{aligned}""",
            r"""\begin{aligned}
r_+ &= M + \sqrt{M^2-a^2},\\
\Omega_{\pm}(r) &= \frac{\pm M^{1/2}}{r^{3/2}\pm aM^{1/2}}
\end{aligned}""",
        ],
        "figures": [
            ("doppler", "图 7  多普勒偏移开启后，吸积盘两侧出现明显的红蓝不对称。", 5.8),
            ("beaming", "图 8  束射增强开启后，高速迎面一侧的亮度显著提高。", 5.8),
        ],
    },
    {
        "title": "九、时间扭曲缩放",
        "paragraphs": [
            "时间扭曲模块位于 physics.ts 的 calculateTimeWarp()。它作为局部时间缩放因子，直接作用于开普勒和 N 体演化，并以 body0 邻近参考点重新采样的方式作用于吸积盘动画速度。靠近中心势阱时，轨道推进与纹理流动都会变慢，因此“近处时间更慢”被转化成了可见的运动节奏差。",
            "理论上，时间膨胀写成固有时与坐标时的关系。静止于 Schwarzschild 势阱中的时钟与绕黑洞做圆轨道运动的时钟并不相同：前者只有引力时间膨胀，后者同时包含引力项与轨道速度的运动学项。弱场近似下，这些效应又会统一到势差和速度平方的同一展开里。",
            "同一个度规因子同时控制时间扭曲和引力红移。某处局部时钟相对无穷远观察者变慢多少，同一处窄谱信号向外传播时红移多少，受的是同一项修正；再把局部轨道速度引入，就会得到“引力项 × 运动学项”的总频移因子。所以这里同时减慢轨道推进和吸积盘动画，让靠近势阱时的节奏变化直接体现在画面上。",
            "实现上，Kepler 模式使用中心势源驱动时间扭曲，N 体模式则以最大质量天体作为参考势源；吸积盘动画速度会在 body0 邻近参考位置重新估算一个 timeWarpFactor，因此轨道推进和盘面流动在视觉上保持同向变化。运行时采用的是有界单调控制函数，只保留“势阱越深、局部速度越高、场景节奏越慢”的物理趋势，而不直接把交互倍率本身写成封闭物理公式。",
        ],
        "formulas": [
            r"""\begin{aligned}
d\tau^2 &= -\frac{1}{c^2}g_{\mu\nu}dx^\mu dx^\nu,\\
\frac{d\tau}{dt} &= \sqrt{1-\frac{2GM}{rc^2}},\qquad
\left.\frac{d\tau}{dt}\right|_{\mathrm{circular}} = \sqrt{1-\frac{3GM}{rc^2}},\qquad
\frac{d\tau}{dt} \approx \sqrt{1-\frac{2GM}{rc^2}-\frac{v^2}{c^2}}
\end{aligned}""",
            r"""\begin{aligned}
\frac{\nu_\infty}{\nu_{\mathrm{local}}} &= \frac{d\tau}{dt},\\
D_{\mathrm{tot}} &= D_{\mathrm{SR}}\frac{d\tau}{dt},\qquad
\frac{\Delta f}{f} \approx \frac{\Delta \Phi}{c^2}
\end{aligned}""",
        ],
        "figures": [
            ("dt", "图 9  时间步长与时间扭曲强度共同改变了轨道和盘面动画节奏。", 5.8),
        ],
    },
    {
        "title": "十、后处理效果：Bloom、Tonemapping 与 Gamma",
        "paragraphs": [
            "后处理链路位于 main.ts 的主帧循环中，真实顺序为场景渲染、可选 MSAA resolve、可选 TAA 混合、亮部提取、Bloom 下采样与上采样、Bloom 合成、ACES 近似色调映射、Gamma 校正、可选 FXAA，以及最终的显示重建与上采样。这个顺序保证了历史帧融合发生在线性 HDR 阶段，Bloom 在高亮信息尚未被压缩前完成扩散，而图像空间边缘修饰与最终上采样都落在显示映射之后。",
            "这个顺序还有一个更本质的理由：Bloom 卷积、MSAA resolve 和 TAA 混合本质上都是对辐亮度做线性组合，只有在线性光空间中它们才保留能量关系；如果先做色调映射或 Gamma，再去卷积和平均，高亮溢出、边缘混合和历史权重都会被非线性压缩扭曲。",
            "Bloom 的理论基础不是“把整张图模糊”，而是近似高亮点扩散函数对局部强辐亮区域的卷积。项目的 brightness pass 先用亮度阈值做硬阈值保色提取，随后用多层 2×2 下采样和上采样构造近似的大核低通，再把结果加回原图。与单次大卷积相比，多尺度金字塔更适合实时图形，也更容易形成柔和稳定的光晕边界。",
            "当前 Bloom 的处理链为“亮部提取 - 多级下采样 - 多级上采样 - 原图合成”。这里的下采样和上采样核心核都是 2×2 平均；当它们在多层分辨率上递推时，等效支撑范围会迅速变大，因此能以较低代价逼近宽核扩散。",
            "色调映射负责把 HDR 辐亮度压缩到显示器可承受的动态范围，Gamma 校正则负责把线性空间中的能量计算转换为更符合显示器响应与人眼感知的输出。当前实现使用 ACES 风格的有理函数拟合；当 tonemappingEnabled 关闭时，这个 pass 只执行纹理拷贝。",
            "完整 ACES 还包含输入变换、可选的 Look Modification、Reference Rendering Transform 与 Output Device Transform 等多个阶段，它解决的不只是“压亮度”，还包括宽色域映射、肩部压缩、暗部抬升和输出设备适配。这里采用的拟合曲线保留了 ACES 在单通道亮度压缩上的典型形状，适合实时链路中的稳定显示映射。"
        ],
        "formulas": [
            r"""\mathcal{B}(aL_1+bL_2)=a\,\mathcal{B}(L_1)+b\,\mathcal{B}(L_2),\qquad T(aL_1+bL_2)\neq a\,T(L_1)+b\,T(L_2)""",
            r"""\begin{aligned}
Y &= 0.2125R + 0.7154G + 0.0721B,\\
C_{\mathrm{bright}}(\bm{x}) &= C(\bm{x})\, H\!\left(Y(\bm{x})-\tau\right),\qquad \tau=1
\end{aligned}""",
            r"""B_\sigma(\bm{x}) = \int_{\Omega} C(\bm{\xi})\, G_\sigma(\bm{x}-\bm{\xi})\, d\bm{\xi}""",
            r"""\begin{aligned}
D_{\ell+1}(\bm{x}) &= \frac14 \sum_{\delta\in\{(-1,-1),(1,-1),(-1,1),(1,1)\}} D_\ell(2\bm{x}+\delta),\\
U_\ell(\bm{x}) &= D_\ell(\bm{x}) + \frac14 \sum_{\delta\in\{(-1,-1),(1,-1),(-1,1),(1,1)\}} U_{\ell+1}(\bm{x}+\delta/2)
\end{aligned}""",
            r"""C_{\mathrm{display}}=\operatorname{ODT}\!\bigl(\operatorname{RRT}(C_{\mathrm{scene}})\bigr)""",
            r"""\operatorname{ACES}(x) = \frac{x(2.51x+0.03)}{x(2.43x+0.59)+0.14}""",
            r"""\begin{aligned}
C_{\mathrm{out}} &=
\begin{cases}
\operatorname{ACES}(C_{\mathrm{hdr}})^{1/\gamma}, & \text{tonemappingEnabled}=1,\\
C_{\mathrm{hdr}}, & \text{tonemappingEnabled}=0
\end{cases}
\end{aligned}""",
        ],
        "figures": [
            ("bloom", "图 10  Bloom 层数较低时，亮部扩散范围更集中。", 5.8),
        ],
    },
    {
        "title": "十一、动态步长与追踪范围优化",
        "paragraphs": [
            "性能优化首先体现在光线步进阶段。shader 并不对所有空间统一采用固定步长，而是根据光线到最近天体表面的距离自适应调整：远离主体时快速跳步，接近黑洞或其他天体表面时自动减小步长，以兼顾吞吐量和局部细节。当前实现的基准步长固定为 0.1，单条光线最多迭代 300 次；只有当最近表面距离足够大时，步长才会放大到 $0.5\\,d_{\\mathrm{surface}}$。这里的距离参照量来自 bodySurfaceRadius，而不是吸积盘本身，因此动态步长首先服务于主体附近的几何与透镜细节。",
            "从数值分析角度看，这一节包含两层不同的优化。第一层是步长控制：局部变化剧烈处缩小步长，局部变化平缓处放大步长，把有限采样预算集中到误差增长更快的区域；项目把“最近表面距离”当作误差代理量，因此实现代价低、实时性好。第二层是方向更新：每一步不只前进，还会在中点位置采样一次透镜加速度，用它修正射线方向，从而在较少步数下保住弯曲轨迹的视觉形状。",
            "从计算量看，步进次数近似等于沿路径把局部弧长除以局部步长后的积分和，因此远场粗采样、近场细采样会直接把预算压到误差增长最快的区域。traceMaxDistance 则相当于给积分区间加上一个数据相关的截断半径：只要光线已经远离所有天体外包络的两倍范围，就可以终止追踪，把剩余贡献交给背景天空盒。",
            "该方案属于“固定基准步长上的显式推进 + 中点校正”方案。它的重点在于用极少额外采样，把有限步数尽可能压到弯曲最剧烈、最容易出锯齿和断裂的区域上。",
            "追踪范围优化则由 traceMaxDistance 完成。主循环会依据相机到各天体表面最远点的距离估计一个全局停止长度，当光线已经远离所有可能产生显著贡献的区域时立即退出。这样可以显著减少背景空域中的无效迭代。",
        ],
        "formulas": [
            r"""h_{n+1}=S\,h_n\left(\frac{\varepsilon}{\lVert e_n\rVert}\right)^{1/(p+1)}""",
            r"""N_{\mathrm{steps}}\approx \int_{0}^{L}\frac{ds}{h(s)}""",
            r"""\begin{aligned}
\bm{a}_{\mathrm{lens}}(\bm{x},\bm{v}) &\approx -\frac{3}{2}\,\frac{h^2\bm{x}}{\lVert \bm{x}\rVert^5},\qquad h^2=\lVert \bm{x}\times \bm{v}\rVert^2,\\
\bm{s}_n &= h_n\hat{\bm{d}}_n,\qquad
\bm{s}_n' = \bm{s}_n + \eta_n\, \bm{a}_{\mathrm{lens}}\!\left(\bm{x}_n+\frac12\bm{s}_n,\bm{s}_0\right),\qquad
\bm{x}_{n+1}=\bm{x}_n+\bm{s}_n'
\end{aligned}""",
            r"""\begin{aligned}
h_{\mathrm{step}} &= \max\!\left(h_0,\, 0.5\, d_{\mathrm{surface}}\right),\\
L_{\mathrm{trace}} &= \max\!\left(L_{\min},\, 2\max_i(d_i + R_i)\right)
\end{aligned}""",
        ],
    },
    {
        "title": "十二、抗锯齿",
        "blocks": [
            {
                "type": "paragraph",
                "text": "抗锯齿的理论底层是采样与重建。连续图像在像素网格上离散化后，只要局部空间频率逼近或超过采样频率的一半，就会出现锯齿、闪烁和时间走样。抗锯齿的核心任务，就是更准确地估计像素面积上的平均辐亮度。"
            },
            {
                "type": "table",
                "caption": "表 3  抗锯齿路径与适用范围。",
                "headers": ["模式", "所在位置", "核心逻辑", "主要取舍"],
                "rows": [
                    ["off", "无额外 pass", "直接输出当前帧", "最清晰但锯齿和闪烁最明显"],
                    ["FXAA", "Tonemapping 之后", "亮度早退、判边、沿边与跨边采样、最终夹紧", "代价低，适合快速抚平高对比边界，但会牺牲部分局部锐度"],
                    ["TAA", "Bloom 与 Tonemapping 之前", "双历史缓冲、同 UV 历史混合、亮度差权重、3×3 邻域夹紧", "更擅长压低时间闪烁，但历史权重过高时会带来拖影"],
                ],
            },
            {
                "type": "formula",
                "latex": r"""\begin{aligned}
\omega_s &> 2\omega_{\max},\\
C_p &= \frac{1}{|A_p|}\int_{A_p} L(x,y)\, dx\, dy,\qquad
C_{\mathrm{MSAA}} = \frac{1}{N}\sum_{i=1}^{N} C(\bm{x}_i)
\end{aligned}""",
            },
            {
                "type": "paragraph",
                "text": "FXAA 的核心控制流程如下："
            },
            {
                "type": "formula",
                "latex": r"""\begin{array}{l}
\text{\textbf{Input: }} C_{\mathrm{center}},\, Y_N,Y_S,Y_E,Y_W,\ \text{quality} \\
\Delta Y \leftarrow Y_{\max}-Y_{\min} \\
\text{\textbf{if }} \Delta Y < \max(0.04,\,0.08Y_{\max}) \text{\textbf{ then return }} C_{\mathrm{center}} \\
\mathbf d \leftarrow \arg\max\!\left(|Y_N-Y_S|,\ |Y_E-Y_W|\right) \\
C_{\parallel} \leftarrow \operatorname{sampleAlong}(C,\mathbf d,\text{quality}) \\
C_{\perp} \leftarrow \operatorname{sampleAcross}(C,\mathbf d,\text{quality}) \\
\alpha \leftarrow \operatorname{subpixelWeight}(Y_N,Y_S,Y_E,Y_W,Y_C) \\
C_{\mathrm{fxaa}} \leftarrow \operatorname{clamp}\!\left((1-\alpha)C_{\parallel}+\alpha C_{\perp}\right)
\end{array}""",
            },
            {
                "type": "paragraph",
                "text": "当前 TAA 直接在同一 UV 位置读取历史帧，并用亮度差控制历史权重，核心混合段如下："
            },
            {
                "type": "code",
                "code": """
float lumaCurrent = luma(current);
float lumaHistory = luma(history);
float weight = 1.0 / (1.0 + abs(lumaCurrent - lumaHistory) * taaFeedback);
weight = clamp(weight, 0.05, 0.95);

vec3 blended = mix(current, history, weight);
blended = clamp(blended, cMin * 0.5, cMax * 1.5);
                """,
            },
            {
                "type": "paragraph",
                "text": "从显示链路上看，TAA 位于 Bloom 与 Tonemapping 之前，FXAA 位于 Tonemapping 之后，因此二者解决的走样类型并不相同：TAA 更擅长压低时间闪烁与细碎边缘噪声，FXAA 更适合在最终 SDR 图像上快速抚平高对比边界；而当 TAA 可用时，MSAA resolve 会先把同一像素内的覆盖样本平均成更稳定的当前帧输入，再进入时域融合。"
            },
            {
                "type": "figure",
                "key": "aa",
                "caption": "图 11  FXAA、TAA 与关闭抗锯齿的边缘平滑差异可以直接观察。",
                "width_inch": 5.8,
            },
        ],
    },
    {
        "title": "十三、显示设置与上采样",
        "blocks": [
            {
                "type": "paragraph",
                "text": "显示设置的第一层杠杆是内部渲染分辨率 $renderScale$。项目先按较低的离屏分辨率完成主要着色，再在显示链末端做重建；$renderScale$ 一旦变化，系统会重新分配整条离屏渲染管线，并同步更新内部分辨率与画布分辨率显示。由于像素着色成本近似正比于内部像素总数，这个参数本质上是在直接控制 fill rate 与采样密度。"
            },
            {
                "type": "table",
                "caption": "表 4  上采样方法的重建思路与使用场景。",
                "headers": ["方法", "核或邻域", "核心逻辑", "结果特征"],
                "rows": [
                    ["Bicubic", "4×4 固定核", "按三次插值核在有限邻域内重建", "整体更平滑，细节恢复较保守"],
                    ["Lanczos2", "有限支撑 sinc 核", "用截断 sinc 做固定核重建", "轮廓更锐，但更容易带出振铃"],
                    ["FSR1", "EASU 12 邻域；RCAS 十字邻域", "先做方向感知重建，再做受约束锐化", "低分辨率下轮廓恢复更明显，且仍保持单帧路径"],
                ],
            },
            {
                "type": "table",
                "caption": "表 5  FSR1 常见工程档位与缩放比例。",
                "headers": ["档位", "输入 / 输出比例", "近似 renderScale", "含义"],
                "rows": [
                    ["Ultra Quality", "1 : 1.3", "约 0.77", "保留较多源分辨率细节"],
                    ["Quality", "1 : 1.5", "约 0.67", "画质与性能折中"],
                    ["Balanced", "1 : 1.7", "约 0.59", "进一步压低着色成本"],
                    ["Performance", "1 : 2.0", "0.50", "更偏向性能，细节恢复压力更大"],
                ],
            },
            {
                "type": "paragraph",
                "text": "从重建理论看，上采样的底层问题是：已知离散采样值，如何估计连续图像在更密网格上的像素平均。理想带限情形对应无限支撑的 sinc 重建，但实时图形必须在有限邻域内完成近似，因此会落到固定核重建、方向自适应重建和后置锐化这几类实际策略上。"
            },
            {
                "type": "formula",
                "latex": r"""\begin{aligned}
f(x,y) &= \sum_{m}\sum_{n} f[m,n]\,\operatorname{sinc}(x-m)\operatorname{sinc}(y-n),\\
W_r &= \operatorname{round}(s\, W_d),\qquad
H_r = \operatorname{round}(s\, H_d),\qquad
N_r = W_rH_r \approx s^2W_dH_d,\\
\Delta t_{\mathrm{cap}} &= \frac{1}{f_{\mathrm{cap}}},\qquad
t_{n+1}^{\mathrm{deadline}} = t_n^{\mathrm{deadline}} + \Delta t_{\mathrm{cap}}
\end{aligned}""",
            },
            {
                "type": "formula",
                "latex": r"""\begin{aligned}
w_{\mathrm{bicubic}}(x) &=
\begin{cases}
1.5|x|^3 - 2.5|x|^2 + 1, & |x|\le 1,\\
-0.5|x|^3 + 2.5|x|^2 - 4|x| + 2, & 1<|x|<2,\\
0, & |x|\ge 2,
\end{cases}\\
w_{\mathrm{Lanczos}}(x;a) &= \operatorname{sinc}(x)\operatorname{sinc}(x/a),\qquad |x|<a
\end{aligned}""",
            },
            {
                "type": "paragraph",
                "text": "FSR1 在当前项目中由 EASU 和 RCAS 两个 pass 串联完成，处理流程如下："
            },
            {
                "type": "formula",
                "latex": r"""\begin{array}{l}
\text{\textbf{EASU Input: }} \{b,c,e,f,g,h,i,j,k,l,n,o\},\ \text{frac} \\
\text{estimate } \mathbf d \text{ from local luma gradients} \\
\text{estimate anisotropic length } \ell \text{ and kernel stretch} \\
\text{for each tap } i:\ w_i \leftarrow W_i(\hat{\mathbf d},\ell,\Delta\mathbf p_i) \\
C_{\mathrm{easu}} \leftarrow \frac{\sum_i w_i C_i}{\sum_i w_i} \\
C_{\mathrm{easu}} \leftarrow \operatorname{clamp}\!\left(C_{\mathrm{easu}},\, C_{\min}^{2\times2},\, C_{\max}^{2\times2}\right)
\end{array}""",
            },
            {
                "type": "formula",
                "latex": r"""\begin{array}{l}
\text{\textbf{RCAS Input: }} B,D,E,F,H,\ s_{\sharp} \\
n_z \leftarrow \operatorname{clamp}\!\left(\frac{|0.25(B_L+D_L+F_L+H_L)-E_L|}{\max(L_{\max}-L_{\min},\varepsilon)},\,0,\,1\right) \\
\lambda \leftarrow \ell\,2^{-s_{\sharp}}\cdot n_z \\
C_{\mathrm{rcas}} \leftarrow \frac{\lambda(B+D+F+H)+E}{4\lambda+1}
\end{array}""",
            },
            {
                "type": "paragraph",
                "text": "由于 FSR1 不使用运动矢量和历史帧，它的本质仍是单帧空间重建而不是时域超分；与此同时，独立的 FPS Cap 则通过维护一个逐帧推进的 deadline 序列整形主循环负载，从而减少简单定时器节流带来的相位漂移和抖动累计。"
            },
            {
                "type": "figure",
                "key": "fsr_off",
                "caption": "图 12  关闭 FSR 时，低渲染倍率下的边缘锯齿更明显。",
                "width_inch": 5.8,
            },
            {
                "type": "figure",
                "key": "fsr_on",
                "caption": "图 13  启用 FSR1 后，低分辨率渲染的轮廓与纹理得到恢复。",
                "width_inch": 5.8,
            },
        ],
    },
    {
        "title": "十四、天空盒管理：双格式支持与 HDR 纹理",
        "blocks": [
            {
                "type": "paragraph",
                "text": "天空盒资源由 resources.ts 自动扫描 public/assets/skybox_*。若目录中只有一个文件，就按全景图 panorama 处理；若存在 6 张图，则按 cubemap 处理。这样既兼容传统立方体天空盒，也兼容一张 equirectangular 全景图，扩展资源时无需改动主业务逻辑。"
            },
            {
                "type": "table",
                "caption": "表 6  天空盒资源组织与加载路径。",
                "headers": ["路径", "资源组织", "采样与上传", "说明"],
                "rows": [
                    ["Cubemap", "同目录 6 张面图", "按面名装配为 TEXTURE_CUBE_MAP", "适合传统六面环境贴图，单面失败时回退为固定像素"],
                    ["Panorama (LDR)", "同目录 1 张全景图", "按经纬映射采样并上传为 TEXTURE_2D", "扩展资源最简单，直接服务背景取样"],
                    ["Panorama (HDR)", "同目录 1 个 .hdr 文件", "先做 Radiance RGBE 解析，再优先上传为 RGBA16F", "若浮点上传失败，再回退到色调映射后的 8-bit LDR 纹理"],
                ],
            },
            {
                "type": "table",
                "caption": "表 7  Cubemap 面轴与文件名对应关系。",
                "headers": ["主轴方向", "采样面", "文件名约定"],
                "rows": [
                    ["+X", "右面", "right"],
                    ["-X", "左面", "left"],
                    ["+Y", "上面", "top"],
                    ["-Y", "下面", "bottom"],
                    ["+Z", "前面", "front"],
                    ["-Z", "后面", "back"],
                ],
            },
            {
                "type": "paragraph",
                "text": "几何采样与 HDR 解码仍可用一组紧凑公式概括："
            },
            {
                "type": "formula",
                "latex": r"""\begin{aligned}
u &= \operatorname{fract}\!\left(\frac12-\frac{\operatorname{atan2}(z,x)}{2\pi}\right),\qquad
v = \operatorname{clamp}\!\left(\frac{\arccos(y)}{\pi},\, \varepsilon,\, 1-\varepsilon\right),\\
\bm{C}_{\mathrm{HDR}} &= 2^{E-136}\,(R,G,B),\qquad
\bm{C}_{\mathrm{ldr}} = \left(\frac{\bm{C}_{\mathrm{HDR}}}{1+\bm{C}_{\mathrm{HDR}}}\right)^{1/2.2}
\end{aligned}""",
            },
            {
                "type": "paragraph",
                "text": "Radiance 头信息解析与失败回退路径如下："
            },
            {
                "type": "code",
                "code": """
const resolutionMatch = line.match(/^([+-])Y\\s+(\\d+)\\s+([+-])X\\s+(\\d+)$/);
const scale = Math.pow(2, exponent - 136);
data[pixelIndex] = scanline[x] * scale;
data[pixelIndex + 1] = scanline[x + width] * scale;
data[pixelIndex + 2] = scanline[x + width * 2] * scale;

if (error !== gl.NO_ERROR) {
  const ldr = createToneMappedFallback(hdr);
  gl.texImage2D(gl.TEXTURE_2D, 0, gl.RGBA, hdr.width, hdr.height, 0, gl.RGBA, gl.UNSIGNED_BYTE, ldr);
}
                """,
            },
            {
                "type": "paragraph",
                "text": "对于黑洞可视化而言，天空盒不是装饰性背景，而是射线追踪终点处的方向辐亮度场。shader 会先让 $rayDir$ 在强透镜场中弯折，再用弯折后的方向去环境图取样；因此背景星场的空间频率、动态范围和颜色分布都会直接影响黑洞阴影边界和吸积盘附近的视觉说服力。"
            },
            {
                "type": "figure",
                "key": "skybox",
                "caption": "图 14  切换天空盒后，背景星场的颜色与空间分布随之变化。",
                "width_inch": 5.8,
            },
        ],
    },
    {
        "title": "十五、视角操控与摄像头手势交互",
        "paragraphs": [
            "视角系统的核心是 getCameraLookBasis()。它根据 yaw、pitch、distance、target 和 roll 生成相机位置以及右、上、前三组正交基向量，随后由 worldToScreenPx() 完成屏幕投影。这样做使 3D 场景观察、轨迹线叠加和屏幕坐标换算都共享同一套相机基，对应的是一个围绕目标点的轨道相机与针孔投影模型。",
            "这套视角基并不是直接从一组固定欧拉角矩阵里整块取出，而是先用观察方向构造前向向量，再用 roll 参考向量与叉乘重新正交化。它等价于在 look-at 约束下做一次简化的 Gram-Schmidt，因此即使前向方向接近世界上方向，也能通过备用参考向量避免基向量退化。",
            "交互层统一建立在 pointer 事件之上，而不是只绑定传统 mouse 事件，并显式关闭默认 touch action，因此桌面端和触屏端共用同一套拖拽与平移逻辑。左键拖拽执行轨道观察，右键拖拽执行 view-aligned 平移，滚轮采用指数缩放调整相机距离，双击恢复默认视角，另有名为正视图和俯视图的两组固定观察预设以及自由观察模式。cameraRoll 则提供绕前向轴的额外转角自由度，使演示镜头不只局限于标准水平地平线。FOV 滑条通过改变投影视野尺度直接影响空间压缩感，因此它不仅是一个显示参数，也是在控制透视强弱与平移灵敏度。",
            "摄像头手势包含本地识别和服务器识别两条链路。本地模式调用 MediaPipe Hands；这类系统的底层通常采用“先检测手掌，再在裁剪区域回归 21 个关键点”的 tracking-by-detection 结构，因为手掌框比整只手姿态更稳定、锚框数量更少，也更适合在手部分遮挡和尺度变化时维持实时性。真实跟踪中还会复用上一帧关键点推导出的 ROI，只在置信度下降时重新触发 palm detector，从而把大部分帧的开销留给关键点回归。服务器模式把 160×120 的 JPEG 帧发送到 /api/detect，并在前端用平滑、死区、超时与并发上限约束输入噪声和链路延迟。",
            "两条链路最终都把“开掌”状态下的 palmX、palmY 映射到 setOrbitPointer()，从而把手势输入统一为对轨道相机状态的连续控制空间。就信息流结构而言，这相当于把离散的视觉识别结果重新参数化为连续控制量，再经过低通滤波和阈值门控送入相机控制器。",
        ],
        "formulas": [
            r"""\bm{c} = \bm{t} + d
\begin{bmatrix}
\cos(\mathrm{pitch})(-\cos(\mathrm{yaw}))\\
\sin(\mathrm{pitch})\\
\cos(\mathrm{pitch})\sin(\mathrm{yaw})
\end{bmatrix}""",
            r"""\begin{aligned}
\mathrm{yaw} &= (p_x-0.5)\, 2\pi,\\
\mathrm{pitch} &= \operatorname{clamp}\!\left((0.5-p_y)\cdot 0.75\pi,\,-0.375\pi,\,0.375\pi\right)
\end{aligned}""",
            r"""\begin{aligned}
f_{\mathrm{view}} &= 2\tan(\mathrm{FOV}/2),\qquad
\Delta \bm{t} = \Delta u\, \bm{u} + \Delta v\, \bm{v},\\
\Delta u &= \frac{\Delta x}{H}\, d\, f_{\mathrm{view}},\qquad
\Delta v = \frac{\Delta y}{H}\, d\, f_{\mathrm{view}},\\
d' &= d\, e^{k\Delta y}
\end{aligned}""",
            r"""\begin{aligned}
x_{\mathrm{ndc}} &= \frac{x_c}{z_c\tan(\mathrm{FOV}/2)},\\
y_{\mathrm{ndc}} &= \frac{y_c}{z_c\tan(\mathrm{FOV}/2)}
\end{aligned}""",
            r"""\begin{aligned}
\bm{p}_t &= (1-\alpha)\bm{p}_{t-1} + \alpha \hat{\bm{p}}_t,\\
\|\hat{\bm{p}}_t-\bm{p}_{t-1}\| < \delta &\Rightarrow \bm{p}_t = \bm{p}_{t-1},\\
\text{isOpenPalm}=1 &\iff \text{fingerCount}\ge 4
\end{aligned}""",
        ],
        "figures": [
            ("camera", "图 15  鼠标拖拽可连续改变轨道视角。", 5.8),
            ("gesture", "图 16  开掌状态下，手掌中心位置会驱动视角旋转。", 5.8),
        ],
    },
    {
        "title": "十六、脚本镜头与参数 JSON 录制 / 回放",
        "paragraphs": [
            "录制回放模块由 RecordingManager 统一管理。系统按“至多每隔约 0.016 s 记一帧”的节奏记录 camera、scene 和关键 render 参数，因此 RecordingFrame 本质上是一个接近 60 FPS 的多组关键状态快照协议。这样的设计使回放重建简单、导入导出格式清晰，也便于把同一实验过程作为可复查数据保留下来。",
            "回放时，main.ts 会直接用录制帧覆盖当前场景、相机和渲染参数；若 renderScale 发生变化，还会重新申请渲染目标，保证离屏分辨率与录制时一致。当前 getPlaybackFrame() 依据回放经过时间查找“时间戳不超过当前 elapsed 的最后一帧”，因此系统行为遵循按时间戳驱动的离散快照重放。界面中的“回放进度”滑条也遵循这一时间语义：它通过改写 playbackStartTime，使当前 wall-clock elapsed 重新对应到目标录制时间点。",
            "JSON 导入采用白名单式校验，检查字段类型、帧数上限、文件大小上限、时间戳单调不减、相机是否落入天体内部，以及旧字段兼容转换如 fsrLike -> fsr1、cameraZoom -> cameraFovDeg。录制帧会先规范化为可校验的 JSON 文本，再通过文件下载、文件导入或 localStorage 保存；同时也遵守 JSON 作为有限数值与对象树交换格式的基本约束，不依赖成员顺序表达语义，也不让 NaN / Infinity 这类超出标准数值域的值进入数据流。这样同一套实验过程既能复查，也能复用。",
        ],
        "formulas": [
            r"""\begin{aligned}
\bm{s}_k &= \bigl(\bm{s}^{\mathrm{camera}}_k,\ \bm{s}^{\mathrm{scene}}_k,\ \bm{s}^{\mathrm{render}}_k\bigr),\\
t_{\mathrm{play}} &= t_{\mathrm{wall}}-t_0,\\
k(t_{\mathrm{play}}) &= \max\{k\mid \tau_k \le t_{\mathrm{play}}\}
\end{aligned}""",
            r"""\begin{aligned}
t_0' &= t_{\mathrm{wall}} - p\,\tau_{\max},\qquad p\in[0,1],\\
\tau_{k+1} &\ge \tau_k,\qquad \|\bm{c}_k-\bm{b}_{i,k}\| > r_{i,k}
\end{aligned}""",
            r"""\begin{aligned}
J &= \operatorname{UTF8}\!\left(\operatorname{Serialize}\!\left(\{\tau_k,\bm{s}_k\}_{k=0}^{N-1}\right)\right),\\
\mathcal{L}_{o}[k] &= J,\qquad o=(\text{scheme},\text{host},\text{port})
\end{aligned}""",
        ],
        "figures": [
            ("recording", "图 17  录制、回放、JSON 导入导出与本地保存集成在同一控制面板中。", 5.8),
        ],
    },
    {
        "title": "十七、Web Audio 氛围音频与场景联动",
        "paragraphs": [
            "ambientAudio.ts 用 Web Audio API 按场景状态实时合成声音。每个天体对应一组主振荡器、八度振荡器和一个作用在 detune 上的低频颤音调制器，再配合全局混响和近距离会遇 burst。这样一来，音频成为场景状态的另一种可感知输出，并与画面演化保持同步。",
            "从系统结构看，Web Audio 本质上是一张有向信号图：振荡器、增益节点、卷积器和参数调制器通过 fan-in / fan-out 连接成多分支音频网络，参数变化则通过 AudioParam 自动化在时间轴上连续逼近目标值。采用这种图式结构后，场景中的质量、速度、距离和时间扭曲都能被实时映射到可听参数。",
            "代码对天体质量采用对数归一化，再把质量映射到基频；相机距离和质量共同决定响度；沿视线方向的径向速度通过双曲正切转换成 pitch bend；速度大小映射到颤音速率和深度；场景越分散、时间扭曲越强，混响比例越高。这一整套映射关系属于典型的参数声化设计，并采用“低频场景分析 + 高频音频平滑更新”的双时间尺度控制结构。",
            "从控制论角度看，模块里大量的 lerp 与 setTargetAtTime 都是在实现一阶低通。连续时间里的指数趋近离散化后会变成 $y_k=(1-\\alpha)y_{k-1}+\\alpha x_k$，这正是它能压住抖动又不至于完全迟钝的原因。两天体在近距离高速接近时，还会触发带短攻击、指数衰减和轻微下滑频的 burst 音，强化“会遇”“擦身”与“拉扯”的事件感；浏览器自动播放策略则通过延迟启用与后续交互恢复来处理，从而兼顾交互体验和平台限制。",
        ],
        "formulas": [
            r"""\begin{aligned}
\Delta c &= 1200\log_2\!\left(\frac{f_{\mathrm{obs}}}{f_{\mathrm{ref}}}\right),\\
f_{\mathrm{vib}}(t) &= f_c + \Delta f\,\sin(2\pi f_m t)
\end{aligned}""",
            r"""\begin{aligned}
L_p &= 20\log_{10}\!\left(\frac{p}{p_0}\right),\qquad
I(r) \propto \frac{1}{r^2},\qquad
A(r)\propto \frac{1}{r}
\end{aligned}""",
            r"""\begin{aligned}
y(t) &= (x*h)(t) = \int_{-\infty}^{+\infty} x(\tau)h(t-\tau)\, d\tau,\\
p(t) &= p_\infty + \bigl(p(t_0)-p_\infty\bigr)e^{-(t-t_0)/\tau}
\end{aligned}""",
            r"""\begin{aligned}
\alpha &= 1-e^{-\Delta t/\tau},\\
y_k &= (1-\alpha)y_{k-1}+\alpha x_k
\end{aligned}""",
            r"""\begin{aligned}
f_{\mathrm{glide}}(t) &= f_0\,e^{-t/\tau_f},\\
g_{\mathrm{burst}}(t) &=
\begin{cases}
g_{\max}\!\left(1-e^{-t/\tau_a}\right), & 0\le t<T_a,\\
g(T_a)\, e^{-(t-T_a)/\tau_r}, & t\ge T_a
\end{cases}
\end{aligned}""",
        ],
        "figures": [
            ("audio", "图 18  音频模块与场景参数联动，可独立启停并调整总体音量。", 5.8),
        ],
    },
    {
        "title": "十八、参数面板配置与布局",
        "blocks": [
            {
                "type": "paragraph",
                "text": "参数面板基于 lil-gui 实现，分为常用操作、场景与物理、画面与特效、录制与回放、音频和天体编辑几大组。画面与特效区又细分为渲染分辨率、抗锯齿、天空盒、黑洞主体、吸积盘、相对论效果和后处理；天体编辑区则显式暴露每个天体的位置、速度、质量、类型、尺寸、发光色、发光强度、吸积盘增益与畸变强度。为了让高频渲染循环、回放系统和 GUI 绑定稳定共存，代码在核心场景状态之外还维护了 params、uiScene、uiView、recordingState、audioState 与 renderState 等状态层。"
            },
            {
                "type": "table",
                "caption": "表 8  参数面板的交互组织原则与界面体现。",
                "headers": ["交互原则", "界面体现", "直接作用"],
                "rows": [
                    ["功能分组", "常用操作、场景与物理、画面与特效、录制、音频、天体编辑分组展开", "缩小单次搜索范围，降低初始浏览负担"],
                    ["渐进披露", "FSR1、TAA、时间扭曲、吸积盘、自旋等高级控件按条件显示", "减少无关控件干扰，避免参数面板过早膨胀"],
                    ["状态可见性", "内部尺寸、画布尺寸、录制进度、音频开关等状态同步展示", "让演示过程中的因果关系和当前模式保持可解释"],
                    ["高频前置", "常用按钮、滑条和视图切换位于上层", "缩短操作链路，降低重复操作成本"],
                ],
            },
            {
                "type": "paragraph",
                "text": "面板采用动态显隐。代码通过 setGuiVisible() 隐藏和显示相关控制，例如只有开启 FSR1 才显示锐化滑条，只有进入 TAA 才显示 MSAA 采样数与 TAA 反馈，只有启用时间扭曲才显示强度、势阱强度和距离参考，只有打开吸积盘、多普勒、束射或自旋时才展开对应的细化参数；录制区还会根据 hasFrames、isRecording、isPlayback 的状态切换按钮与进度条。录制、音频和时间扭曲等分组默认折叠，降低初始界面的信息密度。"
            },
            {
                "type": "paragraph",
                "text": "这套布局遵循的是人机交互里的定性原则：通过分组减少选择负担，通过渐进披露控制信息密度，通过更大的控件和更短的操作链支撑高频调参。参数很多，但不会在初始界面一次性压向使用者；功能很全，但每组边界都较清楚，还能同步展示内部尺寸、画布尺寸、录制进度和音频开关等即时状态。这对课程项目尤其重要，因为它直接决定了演示节奏和可解释性。"
            },
            {
                "type": "figure",
                "key": "display",
                "caption": "图 19  参数面板按功能分组，并根据当前模式动态显示相关控制。",
                "width_inch": 5.8,
            },
        ],
    },
    {
        "title": "十九、项目管理方法：Docker、测试、CI、CD 与 GitHub Pages",
        "paragraphs": [
            "工程采用前后端分层。web/ 负责可视化和交互，server/ 负责可选的服务端手势识别，根目录 docker-compose.yml 把二者编排到同一网络中，前端容器通过 Nginx 反向代理 /api/ 和 /health，并预留 /socket.io/ 转发位置。页面入口围绕 HTTP /api/detect 与 /health 组织，这种“前端统一入口 + 后端内网服务”的结构既降低了跨域复杂度，也让 HTTPS、本地摄像头权限和页面路由更容易统一管理。",
            "前端 Dockerfile 采用多阶段构建，先在 build 阶段产出静态资源，再在 prod 阶段交给 Nginx 提供服务；server/start-server.sh 则优先使用虚拟环境启动服务端，符合“尽量通过虚拟环境或容器避免污染本地环境”的工程要求。默认 Compose 配置优先保证通用环境可运行性，因此前端、后端与代理都可以在统一编排下直接启动并联调。",
            "自动化链路分成 CI 与 Pages CD 两部分。ci.yml 在指向 main 的 Pull Request 上执行前端测试、前端构建和服务端单测；deploy-pages.yml 在 main 分支推送后构建 web/dist 并部署到 GitHub Pages。这样一来，代码合并前可以先通过自动化验证，代码合并后又能快速得到可公开访问的演示页面，形成从开发到展示的连续交付链。",
            "测试范围与项目结构直接对应。前端测试以源码回归断言为主，重点检查相对论效果参数、动态步长、录制参数和交互优先级相关代码是否仍然存在；服务端测试则更接近真实单元测试，重点覆盖非法 JSON、超限图片、速率限制和健康检查语义。部署层面通过 Docker、Nginx、Compose 与 Pages 把本地工程组织成三种交付形态：本地前端开发、Docker 一体联调和 GitHub Pages 静态演示；公开演示页面对应的是静态前端产物，本地识别与外接独立后端都可以围绕这一入口继续扩展。"
        ],
        "figures": [
            ("gh_actions", "图 20  GitHub Actions 工作流记录展示了 CI 与 Pages 部署链路的执行结果。", 5.8),
            ("gh_pages", "图 21  GitHub Pages 已发布公开演示页面，并提供稳定的访问入口。", 5.8),
        ],
    },
    {
        "title": "二十、总结",
        "paragraphs": [
            "这个项目把浏览器端黑洞可视化拆成了几层边界清晰的模块：shader 层负责引力透镜、吸积盘、相对论视效、天空盒和后处理；physics.ts 负责开普勒和 N 体动力学以及时间扭曲近似；camera、handGesture、recordingManager 和 ambientAudio 分别承担视角、手势、可复现实验和参数声化；Docker、CI、CD 与 GitHub Pages 则把实现成果补齐为可交付系统。",
            "项目完成了渲染、交互、录制、音频和部署几个主要部分，代码和界面之间能一一对应。理论部分覆盖黑洞透镜、多天体、显示重建和参数声化等关键依据；实现部分则把这些内容压缩到浏览器里可实时运行的规模。",
            "报告里列出的各模块都能在代码、界面或实验现象中找到对应证据，因此既能交代理论来源，也能说明工程落点。这种写法更符合课程设计报告的要求。",
        ],
    },
]


def set_east_asia_font(run, font_name: str) -> None:
    run.font.name = font_name
    run._element.rPr.rFonts.set(qn("w:eastAsia"), font_name)


def clear_paragraph(paragraph) -> None:
    element = paragraph._element
    for child in list(element):
        element.remove(child)


def add_text_run(
    paragraph,
    text: str,
    *,
    font: str,
    size: float,
    bold: bool = False,
    color: RGBColor | None = None,
):
    run = paragraph.add_run(text)
    set_east_asia_font(run, font)
    run.font.size = Pt(size)
    run.bold = bold
    if color is not None:
        run.font.color.rgb = color
    return run


def format_body_paragraph(paragraph) -> None:
    pf = paragraph.paragraph_format
    pf.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    pf.space_after = Pt(6)
    pf.space_before = Pt(0)
    pf.first_line_indent = Pt(0)
    paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY


def style_existing_paragraph(
    paragraph,
    text: str,
    *,
    font: str,
    size: float,
    bold: bool,
    alignment,
) -> None:
    clear_paragraph(paragraph)
    paragraph.alignment = alignment
    run = add_text_run(paragraph, text, font=font, size=size, bold=bold)
    pf = paragraph.paragraph_format
    pf.space_after = Pt(0)
    pf.space_before = Pt(0)
    pf.line_spacing_rule = (
        WD_LINE_SPACING.ONE_POINT_FIVE if font == BODY_FONT else WD_LINE_SPACING.SINGLE
    )
    set_east_asia_font(run, font)


def add_hyperlink(paragraph, text: str, url: str) -> None:
    part = paragraph.part
    r_id = part.relate_to(url, RELATIONSHIP_TYPE.HYPERLINK, is_external=True)
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), r_id)

    new_run = OxmlElement("w:r")
    r_pr = OxmlElement("w:rPr")

    color = OxmlElement("w:color")
    color.set(qn("w:val"), "0563C1")
    r_pr.append(color)

    underline = OxmlElement("w:u")
    underline.set(qn("w:val"), "single")
    r_pr.append(underline)

    fonts = OxmlElement("w:rFonts")
    fonts.set(qn("w:ascii"), LATIN_FONT)
    fonts.set(qn("w:hAnsi"), LATIN_FONT)
    fonts.set(qn("w:eastAsia"), BODY_FONT)
    r_pr.append(fonts)

    size = OxmlElement("w:sz")
    size.set(qn("w:val"), str(int(12 * 2)))
    r_pr.append(size)
    size_cs = OxmlElement("w:szCs")
    size_cs.set(qn("w:val"), str(int(12 * 2)))
    r_pr.append(size_cs)

    new_run.append(r_pr)
    text_node = OxmlElement("w:t")
    text_node.text = text
    new_run.append(text_node)
    hyperlink.append(new_run)
    paragraph._p.append(hyperlink)


def add_heading(doc: Document, text: str, *, level: int = 1) -> None:
    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
    add_text_run(paragraph, text, font=HEADING_FONT, size=16 if level == 1 else 14, bold=True)
    pf = paragraph.paragraph_format
    pf.space_before = Pt(10)
    pf.space_after = Pt(8)
    pf.line_spacing_rule = WD_LINE_SPACING.SINGLE


def normalize_formula_png(png_path: Path) -> None:
    with Image.open(png_path) as image:
        normalized = image.copy()
    normalized.save(png_path, dpi=(FORMULA_RASTER_DPI, FORMULA_RASTER_DPI))


def prepare_figure_png(source_path: Path, figure_dir: Path) -> Path:
    figure_dir.mkdir(parents=True, exist_ok=True)
    digest = hashlib.sha1(
        f"{source_path.resolve()}::{source_path.stat().st_mtime_ns}::{FIGURE_RASTER_DPI}".encode("utf-8")
    ).hexdigest()[:16]
    png_path = figure_dir / f"{digest}.png"
    if png_path.exists():
        return png_path

    with Image.open(source_path) as image:
        normalized = image.convert("RGBA" if "A" in image.getbands() else "RGB")
        normalized.save(
            png_path,
            format="PNG",
            dpi=(FIGURE_RASTER_DPI, FIGURE_RASTER_DPI),
            optimize=False,
            compress_level=1,
        )
    return png_path


def render_latex_to_png(latex: str, equation_dir: Path, *, inline: bool = False) -> Path:
    equation_dir.mkdir(parents=True, exist_ok=True)
    border_pt = 0 if inline else 1
    font_pt = FORMULA_FONT_PT
    digest = hashlib.sha1(
        f"{'inline' if inline else 'display'}::{border_pt}::{font_pt}::{FORMULA_RASTER_DPI}::{latex}".encode("utf-8")
    ).hexdigest()[:16]
    tex_path = equation_dir / f"{digest}.tex"
    pdf_path = equation_dir / f"{digest}.pdf"
    dvi_path = equation_dir / f"{digest}.dvi"
    png_path = equation_dir / f"{digest}.png"

    if png_path.exists():
        return png_path

    latex_body = rf"\mbox{{\({latex}\)}}" if inline else rf"\mbox{{$\displaystyle {latex}$}}"
    tex_source = textwrap.dedent(
        rf"""
        \documentclass[border={border_pt}pt,{font_pt}pt]{{standalone}}
        \usepackage{{amsmath,amssymb,bm,mathtools}}
        \setlength{{\jot}}{{1pt}}
        \renewcommand{{\arraystretch}}{{0.95}}
        \begin{{document}}
        {latex_body}
        \end{{document}}
        """
    ).strip()
    tex_path.write_text(tex_source, encoding="utf-8")

    pdftoppm_path = PDFTOPPM if Path(PDFTOPPM).exists() else shutil.which("pdftoppm")
    if pdftoppm_path:
        subprocess.run(
            [PDFLATEX, "-interaction=nonstopmode", "-halt-on-error", tex_path.name],
            cwd=equation_dir,
            check=True,
            stdout=subprocess.PIPE,
            stderr=subprocess.STDOUT,
            text=True,
        )
        subprocess.run(
            [pdftoppm_path, "-png", "-r", str(FORMULA_RASTER_DPI), "-singlefile", pdf_path.name, png_path.stem],
            cwd=equation_dir,
            check=True,
            stdout=subprocess.PIPE,
            stderr=subprocess.STDOUT,
            text=True,
        )
        normalize_formula_png(png_path)
        return png_path

    subprocess.run(
        [LATEX, "-interaction=nonstopmode", "-halt-on-error", tex_path.name],
        cwd=equation_dir,
        check=True,
        stdout=subprocess.PIPE,
        stderr=subprocess.STDOUT,
        text=True,
    )
    subprocess.run(
        [DVIPNG, "-D", str(FORMULA_RASTER_DPI), "-T", "tight", "-bg", "Transparent", "-o", png_path.name, dvi_path.name],
        cwd=equation_dir,
        check=True,
        stdout=subprocess.PIPE,
        stderr=subprocess.STDOUT,
        text=True,
    )
    normalize_formula_png(png_path)
    return png_path


def add_formula(doc: Document, latex: str, equation_dir: Path) -> None:
    png_path = render_latex_to_png(latex, equation_dir)
    with Image.open(png_path) as image:
        width_inch = min(6.1, image.width / FORMULA_RASTER_DPI)

    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    pf = paragraph.paragraph_format
    pf.space_before = Pt(0)
    pf.space_after = Pt(4)
    pf.line_spacing_rule = WD_LINE_SPACING.SINGLE
    run = paragraph.add_run()
    run.add_picture(str(png_path), width=Inches(width_inch))


INLINE_MATH_PATTERN = re.compile(r"(\$[^$]+\$)")


def add_inline_formula_run(paragraph, latex: str, equation_dir: Path) -> None:
    png_path = render_latex_to_png(latex, equation_dir, inline=True)
    run = paragraph.add_run()
    run.add_picture(str(png_path))


def add_body(doc: Document, text: str, equation_dir: Path) -> None:
    paragraph = doc.add_paragraph()
    format_body_paragraph(paragraph)
    for chunk in INLINE_MATH_PATTERN.split(text):
        if not chunk:
            continue
        if chunk.startswith("$") and chunk.endswith("$"):
            add_inline_formula_run(paragraph, chunk[1:-1], equation_dir)
        else:
            add_text_run(paragraph, chunk, font=BODY_FONT, size=12)


def set_paragraph_shading(paragraph, fill: str) -> None:
    p_pr = paragraph._element.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), fill)
    p_pr.append(shd)


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_text(
    cell,
    text: str,
    *,
    font: str,
    size: float,
    bold: bool = False,
    alignment=WD_ALIGN_PARAGRAPH.LEFT,
) -> None:
    paragraph = cell.paragraphs[0]
    clear_paragraph(paragraph)
    paragraph.alignment = alignment
    pf = paragraph.paragraph_format
    pf.space_before = Pt(0)
    pf.space_after = Pt(0)
    pf.first_line_indent = Pt(0)
    pf.line_spacing_rule = WD_LINE_SPACING.SINGLE
    lines = text.splitlines() or [text]
    for line_idx, line in enumerate(lines):
        run = add_text_run(paragraph, line, font=font, size=size, bold=bold)
        set_east_asia_font(run, font)
        if line_idx != len(lines) - 1:
            run.add_break()


def add_table_block(
    doc: Document,
    caption: str,
    headers: list[str],
    rows: list[list[str]],
) -> None:
    add_caption(doc, caption)
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = True

    for idx, header in enumerate(headers):
        cell = table.rows[0].cells[idx]
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        set_cell_shading(cell, "D9E2F3")
        set_cell_text(
            cell,
            header,
            font=HEADING_FONT,
            size=10.5,
            bold=True,
            alignment=WD_ALIGN_PARAGRAPH.CENTER,
        )

    for row in rows:
        cells = table.add_row().cells
        for idx, value in enumerate(row):
            cell = cells[idx]
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            set_cell_text(
                cell,
                value,
                font=BODY_FONT,
                size=10.5,
                alignment=WD_ALIGN_PARAGRAPH.LEFT if idx > 0 else WD_ALIGN_PARAGRAPH.CENTER,
            )

    paragraph = doc.add_paragraph()
    paragraph.paragraph_format.space_after = Pt(6)


def add_code_block(doc: Document, code: str) -> None:
    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
    pf = paragraph.paragraph_format
    pf.space_before = Pt(2)
    pf.space_after = Pt(8)
    pf.first_line_indent = Pt(0)
    pf.line_spacing_rule = WD_LINE_SPACING.SINGLE
    set_paragraph_shading(paragraph, "F3F5F7")
    run = add_text_run(paragraph, textwrap.dedent(code).strip("\n"), font=CODE_FONT, size=10.5)
    set_east_asia_font(run, CODE_FONT)


def add_caption(doc: Document, text: str) -> None:
    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    pf = paragraph.paragraph_format
    pf.space_before = Pt(2)
    pf.space_after = Pt(8)
    pf.line_spacing_rule = WD_LINE_SPACING.SINGLE
    add_text_run(paragraph, text, font=BODY_FONT, size=10.5)


def add_figure(doc: Document, path: Path, caption: str, width_inch: float, figure_dir: Path) -> None:
    if not path.exists():
        return
    embedded_path = prepare_figure_png(path, figure_dir)
    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run()
    run.add_picture(str(embedded_path), width=Inches(width_inch))
    add_caption(doc, caption)


def add_opening_resources(doc: Document) -> None:
    paragraph = doc.add_paragraph()
    format_body_paragraph(paragraph)
    add_text_run(paragraph, "GitHub 仓库：", font=BODY_FONT, size=12, bold=True)
    add_hyperlink(
        paragraph,
        "cmzbutqq/MultiMedia-CourseDesign-Insider",
        "https://github.com/cmzbutqq/MultiMedia-CourseDesign-Insider",
    )

    paragraph = doc.add_paragraph()
    format_body_paragraph(paragraph)
    add_text_run(paragraph, "演示页面：", font=BODY_FONT, size=12, bold=True)
    add_hyperlink(
        paragraph,
        "https://cmzbutqq.github.io/MultiMedia-CourseDesign-Insider/",
        "https://cmzbutqq.github.io/MultiMedia-CourseDesign-Insider/",
    )

    paragraph = doc.add_paragraph()
    format_body_paragraph(paragraph)
    add_text_run(paragraph, "演示视频：", font=BODY_FONT, size=12, bold=True)
    add_hyperlink(
        paragraph,
        "https://www.bilibili.com/video/BV1p3NW6DEy1/",
        "https://www.bilibili.com/video/BV1p3NW6DEy1/",
    )

    paragraph = doc.add_paragraph()
    format_body_paragraph(paragraph)
    add_text_run(paragraph, "参考项目：", font=BODY_FONT, size=12, bold=True)
    add_hyperlink(
        paragraph,
        "rossning92/Blackhole",
        "https://github.com/rossning92/Blackhole",
    )


def add_sections(doc: Document, sections: Iterable[dict], equation_dir: Path, figure_dir: Path) -> None:
    for section in sections:
        add_heading(doc, section["title"])
        if section.get("blocks"):
            for block in section["blocks"]:
                block_type = block["type"]
                if block_type == "opening_resources":
                    add_opening_resources(doc)
                elif block_type == "paragraph":
                    add_body(doc, block["text"], equation_dir)
                elif block_type == "formula":
                    add_formula(doc, block["latex"], equation_dir)
                elif block_type == "table":
                    add_table_block(doc, block["caption"], block["headers"], block["rows"])
                elif block_type == "code":
                    add_code_block(doc, block["code"])
                elif block_type == "figure":
                    add_figure(doc, FIGURES[block["key"]], block["caption"], block["width_inch"], figure_dir)
            continue

        if section.get("opening_resources"):
            add_opening_resources(doc)
        for paragraph in section["paragraphs"]:
            add_body(doc, paragraph, equation_dir)
        for formula in section.get("formulas", []):
            add_formula(doc, formula, equation_dir)
        for key, caption, width_inch in section.get("figures", []):
            add_figure(doc, FIGURES[key], caption, width_inch, figure_dir)


def fill_cover(doc: Document) -> None:
    style_existing_paragraph(
        doc.paragraphs[3],
        "多媒体技术课程设计报告书",
        font=HEADING_FONT,
        size=32,
        bold=True,
        alignment=WD_ALIGN_PARAGRAPH.CENTER,
    )
    cover_values = {
        6: TITLE,
        10: "\t\t\t学  院\t\t计算机科学与工程学院",
        11: "\t\t\t专  业\t\t计算机科学与技术",
        12: "\t\t\t组长姓名\t\t陈梦泽",
        13: "\t\t\t组  员\t\t曹丹  陈紫暄  廖晨扬  莫镇嘉",
        14: "\t\t\t指导教师\t\t张艳青",
        15: "\t\t\t课程编号\t\t045101715",
        16: "\t\t\t课程学分\t\t2.0",
        17: "\t\t\t起始日期\t\t2026年3月",
    }

    for index, text in cover_values.items():
        paragraph = doc.paragraphs[index]
        style_existing_paragraph(
            paragraph,
            text,
            font=HEADING_FONT if index == 6 else BODY_FONT,
            size=22 if index == 6 else 15,
            bold=True,
            alignment=(
                WD_ALIGN_PARAGRAPH.CENTER
                if index == 6
                else WD_ALIGN_PARAGRAPH.LEFT
            ),
        )


def prepare_body_start(doc: Document) -> None:
    placeholder = doc.paragraphs[20]
    clear_paragraph(placeholder)
    placeholder.alignment = WD_ALIGN_PARAGRAPH.LEFT
    pf = placeholder.paragraph_format
    pf.space_before = Pt(0)
    pf.space_after = Pt(0)
    pf.line_spacing_rule = WD_LINE_SPACING.SINGLE


def enable_high_fidelity_images(doc: Document) -> None:
    settings = doc.settings.element
    if settings.find(qn("w:doNotCompressPictures")) is None:
        settings.append(OxmlElement("w:doNotCompressPictures"))


def build_report() -> Path:
    doc = Document(str(TEMPLATE_PATH))
    enable_high_fidelity_images(doc)
    fill_cover(doc)
    prepare_body_start(doc)
    with tempfile.TemporaryDirectory(prefix="mm_report_assets_") as tmp_dir:
        tmp_root = Path(tmp_dir)
        add_sections(doc, SECTIONS, tmp_root / "equations", tmp_root / "figures")
        doc.save(str(OUTPUT_PATH))
    return OUTPUT_PATH


if __name__ == "__main__":
    output = build_report()
    print(output)
